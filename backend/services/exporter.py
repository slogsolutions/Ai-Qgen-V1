from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement, ns
from datetime import datetime
import os

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def add_pto_conditional(run):
    fldChar_if_begin = create_element('w:fldChar')
    create_attribute(fldChar_if_begin, 'w:fldCharType', 'begin')
    run._r.append(fldChar_if_begin)
    
    instrText_if1 = create_element('w:instrText')
    create_attribute(instrText_if1, 'xml:space', 'preserve')
    instrText_if1.text = "IF "
    run._r.append(instrText_if1)
    
    fldChar_page_begin = create_element('w:fldChar')
    create_attribute(fldChar_page_begin, 'w:fldCharType', 'begin')
    run._r.append(fldChar_page_begin)
    
    instrText_page = create_element('w:instrText')
    create_attribute(instrText_page, 'xml:space', 'preserve')
    instrText_page.text = "PAGE"
    run._r.append(instrText_page)
    
    fldChar_page_end = create_element('w:fldChar')
    create_attribute(fldChar_page_end, 'w:fldCharType', 'end')
    run._r.append(fldChar_page_end)
    
    instrText_cmp = create_element('w:instrText')
    create_attribute(instrText_cmp, 'xml:space', 'preserve')
    instrText_cmp.text = " < "
    run._r.append(instrText_cmp)
    
    fldChar_nump_begin = create_element('w:fldChar')
    create_attribute(fldChar_nump_begin, 'w:fldCharType', 'begin')
    run._r.append(fldChar_nump_begin)
    
    instrText_nump = create_element('w:instrText')
    create_attribute(instrText_nump, 'xml:space', 'preserve')
    instrText_nump.text = "NUMPAGES"
    run._r.append(instrText_nump)
    
    fldChar_nump_end = create_element('w:fldChar')
    create_attribute(fldChar_nump_end, 'w:fldCharType', 'end')
    run._r.append(fldChar_nump_end)
    
    instrText_cond = create_element('w:instrText')
    create_attribute(instrText_cond, 'xml:space', 'preserve')
    instrText_cond.text = " \"P.T.O.\" \"\" "
    run._r.append(instrText_cond)
    
    fldChar_if_end = create_element('w:fldChar')
    create_attribute(fldChar_if_end, 'w:fldCharType', 'end')
    run._r.append(fldChar_if_end)

def number_to_hindi_word(n: int) -> str:
    words = {1: "एक", 2: "दो", 3: "तीन", 4: "चार", 5: "पाँच", 6: "छह", 7: "सात", 8: "आठ", 9: "नौ", 10: "दस", 11: "ग्यारह", 12: "बारह", 13: "तेरह", 14: "चौदह", 15: "पंद्रह"}
    return words.get(n, str(n))

def number_to_english_word(n: int) -> str:
    words = {1: "one", 2: "two", 3: "three", 4: "four", 5: "five", 6: "six", 7: "seven", 8: "eight", 9: "nine", 10: "ten", 11: "eleven", 12: "twelve", 13: "thirteen", 14: "fourteen", 15: "fifteen"}
    return words.get(n, str(n))

def export_paper_docx(sections_data: dict, section_meta: dict, subject_info: dict, is_answer_key: bool = False) -> str:
    """
    Exports a structured question paper or answer key matching real university exam formatting.
    """
    doc = Document()
    
    # --- Global Style: Set default font ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(2)
    
    # --- Page margins ---
    section_fmt = doc.sections[0]
    section_fmt.top_margin = Inches(0.6)
    section_fmt.bottom_margin = Inches(0.6)
    section_fmt.left_margin = Inches(0.8)
    section_fmt.right_margin = Inches(0.8)
    
    usable_width = 6.9  # 8.5 - 0.8 - 0.8
    
    # --- 1. HEADER BLOCK ---
    title_text = subject_info.get('exam_title', 'EXAMINATION, 2025')
    if is_answer_key:
        title_text = f"ANSWER KEY — {title_text}"
        
    title_par = doc.add_paragraph()
    title_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_par.paragraph_format.space_after = Pt(6)
    title_run = title_par.add_run(title_text)
    title_run.bold = True
    title_run.font.size = Pt(14)

    # Details Block
    details = [
        ("Branch Name", subject_info.get('branch', '')),
        ("Branch Code", subject_info.get('branch_code', '')),
        ("Semester", subject_info.get('sem', '')),
        ("Subject Name", subject_info.get('subject_name', '')),
        ("Subject Code", subject_info.get('subject_code', '')),
    ]
    for label, value in details:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Inches(1.2)
        run = p.add_run(f"{label:>14s}  :  {value}")
        run.font.size = Pt(11)
    
    # Time and M.M. Line (Split)
    time_mm_par = doc.add_paragraph()
    time_mm_par.paragraph_format.tab_stops.add_tab_stop(Inches(usable_width), WD_TAB_ALIGNMENT.RIGHT)
    time_mm_par.paragraph_format.space_before = Pt(6)
    time_mm_par.paragraph_format.space_after = Pt(2)
    r1 = time_mm_par.add_run(f"Time - {subject_info.get('duration', '3:00 Hrs.')}")
    r1.bold = True
    r1.font.size = Pt(11)
    time_mm_par.add_run("\t")
    r2 = time_mm_par.add_run(f"M.M. : {int(subject_info.get('total_marks', 100))}")
    r2.bold = True
    r2.font.size = Pt(11)
    
    # Horizontal rule
    hr = doc.add_paragraph()
    hr.paragraph_format.space_after = Pt(2)
    hr.paragraph_format.space_before = Pt(2)
    hr_run = hr.add_run("_" * 85)
    hr_run.font.size = Pt(8)
    
    # Global Note (Bilingual)
    if not is_answer_key:
        en_note_parts = []
        hi_note_parts = []
        hindi_sec_map = {"A": "क", "B": "ख", "C": "ग", "D": "घ"}
        for sec_title, sec_qs in sections_data.items():
            meta = section_meta.get(sec_title, {})
            attempt = meta.get("attempt_any", len(sec_qs))
            sec_char = sec_title[-1] if sec_title else ""
            h_char = hindi_sec_map.get(sec_char, sec_char)
            en_note_parts.append(f"{attempt} questions from section {sec_char}")
            hi_note_parts.append(f"भाग {h_char} से {number_to_hindi_word(attempt)} प्रश्नों के उत्तर देने हैं")
        
        if len(en_note_parts) > 1:
            en_note_str = "Note : Attempt " + ", ".join(en_note_parts[:-1]) + " and " + en_note_parts[-1] + "."
            hi_note_str = "नोट : " + ", ".join(hi_note_parts[:-1]) + " तथा " + hi_note_parts[-1] + "।"
        else:
            en_note_str = "Note : Attempt " + en_note_parts[0] + "."
            hi_note_str = "नोट : " + hi_note_parts[0] + "।"
        
        np1 = doc.add_paragraph()
        np1.paragraph_format.space_after = Pt(0)
        r = np1.add_run(en_note_str)
        r.bold = True
        r.font.size = Pt(10)
        
        np2 = doc.add_paragraph()
        np2.paragraph_format.space_after = Pt(4)
        r = np2.add_run(hi_note_str)
        r.bold = True
        r.font.size = Pt(10)

    # --- 2. SECTIONS & QUESTIONS ---
    q_counter = 1
    hindi_sec_map = {"A": "क", "B": "ख", "C": "ग", "D": "घ"}
    
    for sec_title, sec_qs in sections_data.items():
        if not sec_qs:
            continue
            
        sec_char = sec_title[-1] if sec_title else ""
        h_char = hindi_sec_map.get(sec_char, sec_char)
        
        # Section Header
        sec_par = doc.add_paragraph()
        sec_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sec_par.paragraph_format.space_before = Pt(10)
        sec_par.paragraph_format.space_after = Pt(2)
        sr = sec_par.add_run(f"SECTION–{sec_char} / भाग–{h_char}")
        sr.bold = True
        sr.underline = True
        sr.font.size = Pt(12)
        
        # Section Note
        if not is_answer_key:
            meta = section_meta.get(sec_title, {})
            attempt_any = meta.get("attempt_any", len(sec_qs))
            marks = meta.get("marks_per_q", 2)
            total = int(attempt_any * marks)
            marks_int = int(marks)
            
            sec_note_par = doc.add_paragraph()
            sec_note_par.paragraph_format.tab_stops.add_tab_stop(Inches(usable_width), WD_TAB_ALIGNMENT.RIGHT)
            sec_note_par.paragraph_format.space_after = Pt(6)
            en_word = number_to_english_word(attempt_any)
            hi_word = number_to_hindi_word(attempt_any)
            sn_run = sec_note_par.add_run(
                f"Note : Attempt any {en_word} questions. / किन्हीं {hi_word} प्रश्नों के उत्तर दीजिये।"
                f"\t{attempt_any}×{marks_int} = {total}"
            )
            sn_run.bold = True
            sn_run.font.size = Pt(10)
            
        # --- Questions ---
        for q_item in sec_qs:
            q_type = q_item.get("q_type", "SA")
            q_en = q_item.get("q", "") or ""
            q_hi = q_item.get("q_hi", "") or ""
            a_en = q_item.get("a", "") or ""
            a_hi = q_item.get("a_hi", "") or ""
            
            # ===== ANSWER KEY =====
            if is_answer_key:
                # English Q + A
                qa_par = doc.add_paragraph()
                qa_par.paragraph_format.space_after = Pt(0)
                qr = qa_par.add_run(f"Q{q_counter}. ")
                qr.bold = True
                qa_par.add_run(q_en)
                
                ans_par = doc.add_paragraph()
                ans_par.paragraph_format.left_indent = Inches(0.4)
                ans_par.paragraph_format.space_after = Pt(0)
                ar = ans_par.add_run("Ans: ")
                ar.bold = True
                ans_par.add_run(a_en if a_en else "—")
                
                # Hindi Q + A (if available)
                if q_hi and q_hi.strip():
                    hi_q_par = doc.add_paragraph()
                    hi_q_par.paragraph_format.left_indent = Inches(0.4)
                    hi_q_par.paragraph_format.space_after = Pt(0)
                    hi_q_par.add_run(q_hi)
                    
                if a_hi and a_hi.strip():
                    hi_a_par = doc.add_paragraph()
                    hi_a_par.paragraph_format.left_indent = Inches(0.4)
                    hi_a_par.paragraph_format.space_after = Pt(0)
                    har = hi_a_par.add_run("उत्तर: ")
                    har.bold = True
                    hi_a_par.add_run(a_hi)
                
                # Small spacer
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_after = Pt(4)
                spacer.paragraph_format.space_before = Pt(0)
                
                q_counter += 1
                continue
            
            # ===== QUESTION PAPER =====
            
            # FIB dots — short and controlled
            FIB_DOTS = "............"
            if q_type == "FIB":
                if "___" in q_en:
                    q_en = q_en.replace("___", FIB_DOTS)
                elif "..." not in q_en:
                    q_en += " " + FIB_DOTS
                if q_hi:
                    if "___" in q_hi:
                        q_hi = q_hi.replace("___", FIB_DOTS)
                    elif "..." not in q_hi:
                        q_hi += " " + FIB_DOTS
                        
            # --- English Question Line ---
            en_par = doc.add_paragraph()
            en_par.paragraph_format.space_after = Pt(0)
            en_par.paragraph_format.left_indent = Inches(0.35)
            en_par.paragraph_format.first_line_indent = Inches(-0.35)
            
            if q_type == "T/F":
                en_par.paragraph_format.tab_stops.add_tab_stop(Inches(usable_width), WD_TAB_ALIGNMENT.RIGHT)
                en_par.add_run(f"{q_counter}.  {q_en}\t(True/False)")
            else:
                en_par.add_run(f"{q_counter}.  {q_en}")
                
            # --- Hindi Question Line ---
            if q_hi and q_hi.strip():
                hi_par = doc.add_paragraph()
                hi_par.paragraph_format.space_after = Pt(2)
                hi_par.paragraph_format.left_indent = Inches(0.55)
                
                if q_type == "T/F":
                    hi_par.paragraph_format.tab_stops.add_tab_stop(Inches(usable_width), WD_TAB_ALIGNMENT.RIGHT)
                    hi_par.add_run(f"{q_hi}\t(सत्य/असत्य)")
                else:
                    hi_par.add_run(q_hi)

            # Options for MCQ
            options = q_item.get("options")
            if options:
                if isinstance(options, dict):
                    for k, v in options.items():
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = Inches(0.7)
                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.space_before = Pt(0)
                        p.add_run(f"{k}) {v}")
                elif isinstance(options, list):
                    for opt in options:
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = Inches(0.7)
                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.space_before = Pt(0)
                        p.add_run(str(opt))
                        
            q_counter += 1
            
    # --- 3. FOOTER BLOCK ---
    footer = section_fmt.footer
    footer_par = footer.paragraphs[0]
    
    tab_stops = footer_par.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(3.25), WD_TAB_ALIGNMENT.CENTER)
    tab_stops.add_tab_stop(Inches(usable_width), WD_TAB_ALIGNMENT.RIGHT)

    footer_par.text = f"{subject_info.get('subject_code', 'SUB')}\t( "
    add_page_number(footer_par.add_run())
    footer_par.add_run(" )\t")
    add_pto_conditional(footer_par.add_run())
    
    # 4. File Management
    os.makedirs("exports", exist_ok=True)
    filename = f"exports/{subject_info.get('subject_code', 'SUB')}_{datetime.now().strftime('%Y%m%d%H%M%S')}"
    if is_answer_key:
        filename += "_ANS"
    filename += ".docx"
    doc.save(filename)
    
    return filename
